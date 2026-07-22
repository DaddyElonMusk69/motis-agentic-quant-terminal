from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import HRFlowable, KeepTogether, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parent
DOCX_PATH = ROOT / "秦浩凯_AI产品经理_优化版.docx"
PDF_PATH = ROOT / "秦浩凯_AI产品经理_优化版.pdf"

NAME = "秦浩凯"
TITLE = "AI 产品经理 / Agent 产品负责人 / AI 数字化项目负责人"
CONTACT = "成都 / 上海 / 杭州  |  13551325719  |  tonyqin58@gmail.com"
GITHUB = (
    "GitHub: github.com/DaddyElonMusk69/motis-agentic-quant-terminal  |  "
    "github.com/DaddyElonMusk69/askcrystal"
)

SUMMARY = (
    "5 年创业及跨境业务 0-1 经验，UC San Diego 计算机科学本科、全英文工作背景。近一年聚焦 AI Agent 产品，"
    "主导 Motis 与 AskCrystal 从场景识别、路线图、工作流和原型设计，到全栈实现、验证与迭代。熟悉 Agent、"
    "RAG、MCP、Tool Calling、Prompt Engineering 与 AI 编程工具，能够连接业务、用户与研发，在高不确定性场景中"
    "推动产品落地，并以可审计流程、用户行为和业务指标驱动迭代。"
)

SECTIONS = [
    (
        "核心项目",
        [
            (
                "Motis - Agentic Quant Research & Trading Platform",
                "创始人 / 产品与技术负责人  |  2025.12 - 至今",
                [
                    "从 0 到 1 定义产品愿景及 V1/V2 路线图，面向量化研究者与个人投资者，将割裂的数据、研究、验证和交易流程重构为 Agent 驱动的一体化工作台。",
                    "梳理“想法→数据→信号→策略训练→Walk-Forward→回测→晋级→实盘”全流程，设计可回测、可复现、可审计的产品契约及 Agent/系统边界，降低研究到生产的行为漂移。",
                    "以真实研究任务持续 dogfooding，围绕任务完成、迭代效率、可追溯性与风险控制优化体验；用 Codex、Claude Code 和多 Agent 工作流完成需求拆解、交互验证、全栈开发与测试闭环。",
                ],
            ),
            (
                "AskCrystal - AI 电商导购与运营平台",
                "创始人 / 产品与技术负责人  |  2025.12 - 至今",
                [
                    "从 0 到 1 设计 AI 导购与店铺运营一体化产品，打通前台用户决策与后台商品、内容和运营工作流。",
                    "基于 Dify、RAG 与长期记忆构建跨会话个性化导购，按用户意图与生命周期分层；围绕问答质量、转化路径和反馈持续优化推荐体验。",
                    "将商品上架、分类、内容与运营数据管理沉淀为可复用 Skill / Workflow，验证 AI 同时改善交易体验与运营效率的闭环。",
                ],
            ),
        ],
    ),
    (
        "工作经历",
        [
            (
                "香港 Neox Technologies Ltd.",
                "跨境电商合伙人 / 业务产品负责人  |  2021.08 - 至今",
                [
                    "从 0 到 1 搭建欧美跨境业务，负责用户研究、选品与定位、Shopify 产品体验、增长、合规及跨团队交付；与海外客户、合作伙伴及平台保持全英文沟通。",
                    "建立 Meta / Google / TikTok 增长体系，以 ROAS、CAC、AOV 和转化漏斗进行周度复盘及 A/B 实验；3 个月营收增长近 10 倍，客单价约 150 美元，年 GMV 超 1,000 万美元。",
                    "梳理订单、支付、物流、客服及运营协作链路，引入 RPA 与 SOP 完成关键流程数字化，明确责任人、交付节点和异常处理；协调投放、运营、技术及供应链持续上线与迭代。",
                ],
            )
        ],
    ),
]

SKILLS = [
    ("AI 产品", "Agent、RAG、MCP、Tool Calling、Prompt Engineering、LangGraph、Dify、Codex、Claude Code、Harness Engineering、多 Agent 协作"),
    ("产品交付", "产品路线图、PRD、用户访谈、流程与信息架构、AI 辅助原型、需求优先级、里程碑 / 风险 / UAT、跨团队协作"),
    ("数据增长", "指标体系、漏斗分析、A/B 测试、用户分层、转化与留存、ROI / ROAS、海外增长"),
    ("技术与语言", "Python、TypeScript / React、FastAPI、PostgreSQL、Git、API、Shopify、RPA；中文母语、英语可作为工作语言"),
]


def set_cell_shading(element, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    element.get_or_add_pPr().append(shd)


def set_run_font(run, name="Arial", east_asia="STHeiti", size=9.2, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def compact_paragraph(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_docx_section_heading(doc, text):
    p = doc.add_paragraph()
    compact_paragraph(p, before=2.5, after=1.5)
    p.paragraph_format.keep_with_next = True
    set_cell_shading(p._element, "E8F0F2")
    r = p.add_run(f"  {text}")
    set_run_font(r, size=10.5, bold=True, color=(17, 74, 88))
    return p


def add_docx_role(doc, title, meta, bullets):
    p = doc.add_paragraph()
    compact_paragraph(p, before=1.3, after=0.5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, size=9.3, bold=True, color=(22, 59, 70))
    r = p.add_run(f"  |  {meta}")
    set_run_font(r, size=8.5, color=(80, 80, 80))
    for bullet in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        compact_paragraph(bp, after=0.3, line=0.95)
        bp.paragraph_format.left_indent = Cm(0.42)
        bp.paragraph_format.first_line_indent = Cm(-0.22)
        for run in bp.runs:
            set_run_font(run, size=8.45)
        if not bp.runs:
            set_run_font(bp.add_run(bullet), size=8.45)
        else:
            bp.runs[0].text = bullet


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.7)
    section.bottom_margin = Cm(0.65)
    section.left_margin = Cm(0.9)
    section.right_margin = Cm(0.9)
    section.header_distance = Cm(0.25)
    section.footer_distance = Cm(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(9.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact_paragraph(p, after=0)
    set_run_font(p.add_run(NAME), size=18, bold=True, color=(13, 55, 65))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compact_paragraph(p, after=0.3)
    set_run_font(p.add_run(TITLE), size=10.5, bold=True, color=(0, 133, 145))

    for line, size in ((CONTACT, 8.5), (GITHUB, 7.7)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compact_paragraph(p, after=0.2)
        set_run_font(p.add_run(line), size=size, color=(70, 70, 70))

    add_docx_section_heading(doc, "职业概述")
    p = doc.add_paragraph()
    compact_paragraph(p, after=0.5, line=0.96)
    set_run_font(p.add_run(SUMMARY), size=9.0)

    for heading, roles in SECTIONS:
        add_docx_section_heading(doc, heading)
        for role in roles:
            add_docx_role(doc, *role)

    add_docx_section_heading(doc, "核心能力")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        compact_paragraph(p, after=0.2, line=0.94)
        set_run_font(p.add_run(f"{label}："), size=8.6, bold=True, color=(17, 74, 88))
        set_run_font(p.add_run(value), size=8.6)

    add_docx_section_heading(doc, "教育背景")
    p = doc.add_paragraph()
    compact_paragraph(p, after=0)
    set_run_font(p.add_run("University of California San Diego"), size=8.8, bold=True, color=(22, 59, 70))
    set_run_font(p.add_run("  |  计算机科学 学士  |  2015 - 2019"), size=8.5)

    doc.save(DOCX_PATH)


def build_pdf():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=7 * mm,
        bottomMargin=6 * mm,
        title="秦浩凯 - AI 产品经理简历",
        author="秦浩凯",
    )
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "CN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=9.5,
        leading=12.5,
        textColor=colors.HexColor("#202529"),
        spaceAfter=1.2,
        wordWrap="CJK",
    )
    name_style = ParagraphStyle(
        "Name", parent=base, fontSize=20, leading=22, alignment=TA_CENTER,
        textColor=colors.HexColor("#0D3741"), spaceAfter=1.2,
    )
    title_style = ParagraphStyle(
        "Title", parent=base, fontSize=11.4, leading=13.5, alignment=TA_CENTER,
        textColor=colors.HexColor("#008591"), spaceAfter=0.8,
    )
    contact_style = ParagraphStyle(
        "Contact", parent=base, fontSize=8.4, leading=10, alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"), spaceAfter=0.35,
    )
    section_style = ParagraphStyle(
        "Section", parent=base, fontSize=10.7, leading=13,
        textColor=colors.HexColor("#114A58"), backColor=colors.HexColor("#E8F0F2"),
        borderPadding=(1.4, 3.2, 1.4, 3.2), spaceBefore=3.2, spaceAfter=1.8,
    )
    role_style = ParagraphStyle(
        "Role", parent=base, fontSize=9.6, leading=12,
        textColor=colors.HexColor("#163B46"), spaceBefore=1.7, spaceAfter=0.7,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=base, fontSize=8.9, leading=11.5,
        leftIndent=9.0, firstLineIndent=-5.5, bulletIndent=2.5, spaceAfter=0.75,
    )
    skill_style = ParagraphStyle(
        "Skill", parent=base, fontSize=8.6, leading=10.8, spaceAfter=0.55,
    )

    story = [
        Paragraph(NAME, name_style),
        Paragraph(TITLE, title_style),
        Paragraph(CONTACT, contact_style),
        Paragraph(GITHUB, contact_style),
        Paragraph("职业概述", section_style),
        Paragraph(SUMMARY, base),
    ]

    for heading, roles in SECTIONS:
        story.append(Paragraph(heading, section_style))
        for title, meta, bullets in roles:
            story.append(Paragraph(f"<b>{title}</b>  |  <font color='#555555'>{meta}</font>", role_style))
            for bullet in bullets:
                story.append(Paragraph(bullet, bullet_style, bulletText="-"))

    story.append(Paragraph("核心能力", section_style))
    for label, value in SKILLS:
        story.append(Paragraph(f"<b>{label}：</b>{value}", skill_style))

    story.append(Paragraph("教育背景", section_style))
    story.append(Paragraph("<b>University of California San Diego</b>  |  计算机科学 学士  |  2015 - 2019", base))
    doc.build(story)


if __name__ == "__main__":
    ROOT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    page_count = len(PdfReader(str(PDF_PATH)).pages)
    print(f"created: {DOCX_PATH}")
    print(f"created: {PDF_PATH}")
    print(f"pdf_pages: {page_count}")
